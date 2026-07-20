"""Generate Chapters II and III for Selva Booking — UNAMAD / ApMo style."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Capitulo_II_y_III_Selva_Booking.docx"
REF = ROOT / "docs" / "ApMo_Proyecto_HOTEL.docx"


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)


def heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Paragraph")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Caption")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)


def fuente(doc: Document) -> None:
    p = doc.add_paragraph("Fuente: Elaboración propia")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val
            for run in t.rows[ri].cells[ci].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    set_normal_style(doc)
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3)
        s.right_margin = Cm(2.5)

    # ── CAPÍTULO II ─────────────────────────────────────────────
    heading(doc, "CAPÍTULO II: MARCO TEÓRICO", 1)

    heading(doc, "2.1. Antecedentes de estudio", 2)
    para(
        doc,
        "Se revisó el uso de aplicaciones móviles y plataformas en línea para la reserva de "
        "alojamientos turísticos en destinos de naturaleza. Estudios sobre transformación digital "
        "del turismo (UNWTO, 2023) y guías de arquitectura Android (Google, 2024) muestran que "
        "centralizar la oferta hotelera en una sola app mejora la comparación de precios y reduce "
        "el tiempo de búsqueda del viajero. En Madre de Dios, plataformas globales como Booking.com "
        "no siempre incluyen eco-lodges locales, por lo que muchos turistas dependen de páginas "
        "web dispersas. La propuesta de Selva Booking se alinea con este enfoque: una app Android "
        "nativa con Firebase que unifica hoteles urbanos y lodges selváticos, con panel de "
        "administración para que los negocios locales gestionen habitaciones y reservas.",
    )

    heading(doc, "2.2. Marco teórico", 2)

    heading(doc, "2.2.1. Desarrollo de aplicaciones móviles Android", 3)
    para(
        doc,
        "Una aplicación móvil Android es un programa instalado en smartphones que utiliza el "
        "sistema operativo de Google y su kit de desarrollo (SDK) para ofrecer servicios al usuario "
        "de forma nativa. Kotlin es el lenguaje recomendado por Google por ser conciso, seguro "
        "y compatible con las bibliotecas Jetpack. Selva Booking se desarrolla con Kotlin, "
        "Jetpack Compose para las pantallas y Navigation Compose para moverse entre login, "
        "búsqueda, reservas y panel administrativo.",
    )

    heading(doc, "2.2.2. Firebase y servicios en la nube", 3)
    para(
        doc,
        "Firebase es una plataforma de Google que ofrece servicios en la nube —autenticación, "
        "base de datos y almacenamiento de archivos— sin necesidad de administrar servidores "
        "propios. En Selva Booking, Firebase Authentication gestiona el registro e inicio de "
        "sesión; Cloud Firestore guarda en tiempo real usuarios, hoteles, habitaciones y reservas; "
        "y Firebase Storage almacena las fotos de hoteles, habitaciones y perfiles. Esto permite "
        "que los datos se sincronicen entre dispositivos de forma rápida y segura.",
    )

    heading(doc, "2.2.3. Sistemas de reservas hoteleras y ecoturismo", 3)
    para(
        doc,
        "Un sistema de reservas hoteleras permite al turista consultar disponibilidad, comparar "
        "precios y confirmar una estadía en un establecimiento. En el ecoturismo amazónico, "
        "los alojamientos suelen ser eco-lodges y cabañas con capacidad limitada, distintos "
        "a los hoteles urbanos convencionales. Selva Booking modela hoteles, habitaciones y "
        "reservas con estados como Pendiente, Confirmada y Cancelada, e incluye filtros por "
        "ciudad, precio y valoración para facilitar la elección del viajero en Puerto Maldonado "
        "y Tambopata.",
    )

    heading(doc, "2.3. Definición de términos", 2)
    para(doc, "Se definen los acrónimos y conceptos técnicos más relevantes del proyecto Selva Booking.")
    caption(doc, "Tabla 1: Glosario de términos")
    table(
        doc,
        ["Término", "Definición"],
        [
            ("Android", "Sistema operativo móvil donde se ejecuta Selva Booking."),
            ("API", "Interfaz que permite la comunicación entre componentes de software."),
            ("Backend", "Servicios en la nube que procesan y almacenan datos (Firebase en Selva Booking)."),
            ("Cliente", "Usuario turista que busca, compara y reserva alojamientos."),
            ("Cloud Firestore", "Base de datos NoSQL de Firebase usada para hoteles, habitaciones y reservas."),
            ("Eco-lodge", "Alojamiento turístico de bajo impacto ambiental en zona selvática."),
            ("Firebase Authentication", "Servicio de registro e inicio de sesión con correo y contraseña."),
            ("Firebase Storage", "Almacenamiento en la nube de imágenes de hoteles y perfiles."),
            ("Frontend", "Interfaz visible con la que interactúa el usuario (pantallas Compose)."),
            ("Jetpack Compose", "Framework declarativo de Google para diseñar la interfaz de la app."),
            ("Kotlin", "Lenguaje principal de desarrollo Android en Selva Booking."),
            ("MVVM", "Patrón que separa datos, lógica (ViewModel) e interfaz (View)."),
            ("Administrador", "Usuario que gestiona hoteles, habitaciones y reservas desde el panel admin."),
            ("Reserva", "Registro de una estadía con fechas, habitación, huéspedes y precio total."),
            ("UML", "Lenguaje unificado para diagramar casos de uso, clases y secuencias del sistema."),
            ("XP", "Metodología ágil de Programación Extrema usada en el desarrollo del proyecto."),
        ],
    )
    fuente(doc)

    doc.add_page_break()

    # ── CAPÍTULO III ────────────────────────────────────────────
    heading(doc, "CAPÍTULO III: METODOLOGÍA DE LA INVESTIGACIÓN", 1)

    heading(doc, "3.1. Tipo de estudio", 2)
    para(
        doc,
        "La investigación que sustenta el desarrollo de Selva Booking es de tipo aplicada, porque "
        "parte de conocimientos ya existentes en desarrollo móvil y los pone al servicio de una "
        "necesidad real: centralizar la reserva de alojamientos en Madre de Dios.",
    )
    para(
        doc,
        "Según su alcance, es descriptiva-proyectiva, porque describe cómo funciona la aplicación "
        "y propone una solución tecnológica concreta. En cuanto al diseño, es no experimental, "
        "ya que no se manipulan variables en un laboratorio, sino que se construye y evalúa un "
        "producto software. Finalmente, tiene un enfoque mixto: cualitativo al validar usabilidad "
        "con usuarios, y cuantitativo al registrar tiempos de respuesta y cantidad de reservas "
        "de prueba en Firebase.",
    )

    heading(doc, "3.2. Población", 2)

    heading(doc, "3.2.1. Población del modelo", 3)
    para(
        doc,
        "La población del modelo de datos comprende todas las entidades que Selva Booking gestiona "
        "en Firebase: usuarios registrados (clientes y administradores), hoteles y eco-lodges de "
        "Madre de Dios, habitaciones asociadas a cada establecimiento y reservas vinculadas a "
        "usuario, hotel y habitación. Para pruebas se utilizó información de ejemplo con "
        "establecimientos en Puerto Maldonado y Tambopata, con precios, servicios e imágenes "
        "representativas del ecoturismo regional.",
    )

    heading(doc, "3.2.2. Población de usuarios y muestra de validación", 3)
    para(
        doc,
        "La población objetivo de la aplicación son turistas y administradores de alojamientos "
        "mayores de 18 años que cuenten con un dispositivo Android. Para la fase de pruebas de "
        "usabilidad, se trabajará con una muestra de 10 personas de la ciudad de Puerto Maldonado, "
        "seleccionadas de forma no probabilística por conveniencia, buscando incluir clientes que "
        "simulen reservas y al menos dos administradores que prueben el panel de gestión de "
        "hoteles y reservas.",
    )

    heading(doc, "3.3. Metodología de desarrollo", 2)
    para(
        doc,
        "Para construir el aplicativo Selva Booking decidimos trabajar con la metodología ágil "
        "Programación Extrema (XP - Extreme Programming). La razón principal es que nuestro equipo "
        "está formado por dos integrantes, los requisitos del proyecto pueden ir cambiando a medida "
        "que probemos el aplicativo con posibles usuarios, y necesitamos sacar versiones que "
        "funcionen en poco tiempo para verificar si la búsqueda, la reserva y el panel admin "
        "responden correctamente.",
    )
    para(
        doc,
        "XP plantea entregar software de calidad mediante ciclos cortos de trabajo, realizar "
        "pruebas seguido y mantener una buena comunicación dentro del equipo. Estas características "
        "encajan con lo que buscamos en Selva Booking, ya que las pantallas y reglas de negocio "
        "—como filtros de búsqueda, cálculo de noches y estados de reserva— no se van a perfeccionar "
        "de un solo intento, sino que tenemos que ir ajustando el aplicativo poco a poco hasta que "
        "el flujo completo sea confiable.",
    )

    heading(doc, "3.3.1. Fases de la metodología XP", 3)
    para(
        doc,
        "El desarrollo del proyecto se organizó en las cinco etapas que propone XP, las cuales "
        "se repiten de manera iterativa hasta llegar a la versión final del aplicativo.",
    )
    caption(doc, "Tabla 2. Fases de la metodología XP aplicadas a Selva Booking")
    table(
        doc,
        ["Fase", "Descripción general", "Aplicación en Selva Booking"],
        [
            (
                "Planificación",
                "Se definen las historias de usuario, los valores, los criterios de aceptación y el orden de prioridad del trabajo.",
                "Se elaboraron historias de usuario para registro, búsqueda de hoteles, reserva, pago simulado, panel admin y gestión de habitaciones.",
            ),
            (
                "Diseño",
                "Se busca un diseño simple del sistema, apoyándose en tarjetas CRC y prototipos.",
                "Se diseñaron las 24 pantallas principales (wireframes SVG/Figma), la paleta de colores Selva y los diagramas UML del sistema.",
            ),
            (
                "Codificación",
                "Se realiza la programación del producto y se aplica rediseño cuando es necesario.",
                "Se programaron los módulos en Kotlin con Compose, ViewModels, repositorios e integración con Firebase Auth, Firestore y Storage.",
            ),
            (
                "Pruebas",
                "Se aplican pruebas unitarias y refactorización continua para validar el funcionamiento.",
                "Se probaron los módulos en emulador y dispositivo Android real: login, filtros, reserva, cancelación y CRUD administrativo.",
            ),
            (
                "Lanzamiento",
                "Se entrega un incremento del software listo para ser usado.",
                "Se liberó una versión funcional de Selva Booking probada con usuarios de conveniencia en Puerto Maldonado.",
            ),
        ],
    )
    fuente(doc)

    heading(doc, "3.4. Recursos", 2)

    heading(doc, "3.4.1. Recursos humanos", 3)
    para(
        doc,
        "El proyecto es desarrollado por estudiantes de la Escuela Profesional de Ingeniería de "
        "Sistemas e Informática de la UNAMAD, con la asesoría del Mg. Tineo Vilchez, Francisco Javier. "
        "Cada integrante asumió un rol según sus fortalezas:",
    )
    bullet(doc, "Condori Sahuarico, Liz Yeiza — Diseño de interfaces (UI/UX), wireframes y pruebas de usabilidad.")
    bullet(doc, "Uceda Jallurana, Renzo Jesus — Desarrollo Android, integración Firebase, módulos de reserva y panel administrativo.")

    heading(doc, "3.4.2. Recursos tecnológicos", 3)
    para(
        doc,
        "Para el desarrollo de Selva Booking se utilizan herramientas de libre acceso o con planes "
        "gratuitos suficientes para las necesidades del proyecto:",
    )
    caption(doc, "Tabla 3: Especificaciones del entorno de desarrollo y tecnologías aplicadas")
    table(
        doc,
        ["Software", "Versión / Detalle", "Rol"],
        [
            ("Android Studio", "Iguana / última estable", "IDE principal de la app Android"),
            ("Kotlin", "1.9+", "Lenguaje de programación de la app"),
            ("Jetpack Compose", "BOM 2024", "Framework de UI declarativa"),
            ("Firebase BOM", "33.x", "Auth, Firestore y Storage"),
            ("Coil", "2.x", "Carga de imágenes de hoteles y perfiles"),
            ("Navigation Compose", "2.x", "Navegación entre pantallas"),
            ("Python", "3.10+", "Generación de wireframes SVG y documentación"),
            ("Figma", "Plan gratuito", "Prototipado visual de interfaces"),
        ],
    )
    fuente(doc)

    heading(doc, "Presupuesto", 2)
    para(doc, "La Tabla 4 presenta la estructura de costos y el presupuesto general requerido para el desarrollo e implementación del sistema.")
    caption(doc, "Tabla 4: Presupuesto estimado del proyecto")
    table(
        doc,
        ["Categoría", "Recurso", "Costo"],
        [
            ("Software", "Android Studio", "S/. 0"),
            ("Software", "Kotlin + Jetpack Compose", "S/. 0"),
            ("Software", "Figma (plan gratuito)", "S/. 0"),
            ("Servicios en la nube", "Firebase - Plan Spark", "S/. 0"),
            ("Hardware", "Computadoras (personales)", "S/. 0"),
            ("Hardware", "Smartphone Android para pruebas", "S/. 0"),
            ("Otros", "Internet (4 meses)", "S/. 400"),
            ("Contingencia", "Reserva", "S/. 48"),
            ("TOTAL ESTIMADO", "TOTAL ESTIMADO", "S/. 448"),
        ],
    )
    fuente(doc)

    heading(doc, "3.5. Cronograma", 2)
    para(
        doc,
        "El desarrollo de Selva Booking se organiza en seis entregables, los cuales marcan el "
        "avance del proyecto a lo largo del semestre; cada entregable representa una etapa concreta "
        "de construcción del sistema.",
    )
    caption(doc, "Tabla 5. Cronograma")
    table(
        doc,
        ["ENTREGABLE", "CONTENIDO", "PERÍODO"],
        [
            ("1er", "Nombre del proyecto, descripción del proyecto e indicación de la metodología de desarrollo", "Semana 1"),
            ("2do", "Wireframe del aplicativo", "Semana 3"),
            ("3er", "Diagramación UML", "Semana 5"),
            ("4to", "Desarrollo de interfaces y código fuente (primera parte)", "Semana 7"),
            ("5to", "Desarrollo de interfaces y código fuente (segunda parte)", "Semana 10"),
            ("6to", "Presentación final del proyecto con documentación completa y programa funcional", "Semana 13"),
        ],
    )
    fuente(doc)

    return doc


def patch_apmo_chapters(doc_path: Path, new_doc: Document) -> None:
    """Replace Cap II–III body in ApMo reference file with generated content."""
    if not doc_path.exists():
        return
    ref = Document(doc_path)
    start, end = None, None
    for i, p in enumerate(ref.paragraphs):
        t = p.text.strip().upper()
        if start is None and "MARCO TE" in t and "RICO" in t:
            start = i
        if start is not None and end is None and t.startswith("WIREFRAME"):
            end = i
            break
    if start is None or end is None:
        print("No se encontró rango Cap II–III en ApMo; solo se generó archivo independiente.")
        return

    # Remove paragraphs from end-1 down to start (keep WIREFRAME)
    for i in range(end - 1, start - 1, -1):
        el = ref.paragraphs[i]._element
        el.getparent().remove(el)

    # Remove tables that belonged to Cap II–III (first 5 tables in ApMo are glossary+methodology)
    while len(ref.tables) > 0:
        tbl = ref.tables[0]._element
        tbl.getparent().remove(tbl)
        if len(ref.tables) == 0:
            break
        # stop if no more tables at beginning - ApMo has 5 tables in cap II-III
        if len([t for t in ref.tables]) == 0:
            break
    # Remove all 5 methodology tables from ref (they appear before wireframe figures)
    for _ in range(5):
        if ref.tables:
            ref.tables[0]._element.getparent().remove(ref.tables[0]._element)

    anchor = ref.paragraphs[start]._element
    parent = anchor.getparent()
    idx = list(parent).index(anchor)

    # Insert new paragraphs before WIREFRAME
    from docx.oxml import OxmlElement
    from copy import deepcopy

    for block in new_doc.element.body:
        parent.insert(idx, deepcopy(block))
        idx += 1

    ref.save(doc_path)
    print(f"Actualizado: {doc_path}")


def main() -> None:
    doc = build()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Documento generado: {OUT}")
    try:
        patch_apmo_chapters(REF, doc)
    except Exception as e:
        print(f"Aviso: no se pudo parchear ApMo ({e})")


if __name__ == "__main__":
    main()
